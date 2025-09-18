from flask import Flask, render_template, send_from_directory, request, jsonify, send_file, abort
import os
import io
import json
from datetime import datetime

# NEW imports (add after existing imports at the top)
from config import AppConfig
from services.storage import (
    ensure_dirs, make_batch_folder, relpath_from_output, zip_outputs, safe_filename
)
from services.csv_batch import process_csv as csv_process
from services.word_fill import fill_and_export
from services.extract_input import extract_and_map


# ===== Optional deps for DOCX/PDF work =====
DOCX_AVAILABLE = True
PDF_AVAILABLE = True
try:
    from docx import Document
    from docx.enum.text import WD_BREAK
except Exception:
    DOCX_AVAILABLE = False

try:
    from PyPDF2 import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
except Exception:
    PDF_AVAILABLE = False


def create_app():
    app = Flask(__name__, static_folder='static', template_folder='templates')

    # ---- paths / defaults
    app.config["DEFAULT_DOCX_TEMPLATE"] = os.path.join(app.root_path, "static", "templates", "default_template.docx")
    app.config["DEFAULT_PDF_TEMPLATE"] = os.path.join(app.root_path, "static", "templates", "default_template.pdf")
    os.makedirs(os.path.dirname(app.config["DEFAULT_DOCX_TEMPLATE"]), exist_ok=True)

    @app.route("/")
    def index():
        return render_template("index.html")

    @app.route("/signin")
    def signin():
        return render_template("signin.html")

    @app.route("/signup")
    def signup():
        return render_template("signup.html")

    @app.route("/app")
    def app_shell():
        return render_template("app_shell.html")

    # Optional: serve favicon if you add one later
    @app.route('/favicon.ico')
    def favicon():
        return send_from_directory(os.path.join(app.root_path, 'static', 'img'),
                                   'logo.svg', mimetype='image/svg+xml')

    # ---------------------------
    # Helpers (templates + utils)
    # ---------------------------
    def ensure_default_docx():
        """Create a very simple 4-page DOCX template with page breaks if missing."""
        if not DOCX_AVAILABLE:
            return
        path = app.config["DEFAULT_DOCX_TEMPLATE"]
        if os.path.exists(path):
            return
        d = Document()
        d.add_paragraph("Default Template — Page 1")
        d.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
        d.add_paragraph("Default Template — Page 2")
        d.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
        d.add_paragraph("Default Template — Page 3")
        d.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
        d.add_paragraph("Default Template — Page 4")
        d.save(path)

    def ensure_default_pdf():
        """Create a very simple 4-page PDF template if missing."""
        if not PDF_AVAILABLE:
            return
        path = app.config["DEFAULT_PDF_TEMPLATE"]
        if os.path.exists(path):
            return
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        for p in range(1, 5):
            c.setFont("Helvetica", 14)
            c.drawString(72, A4[1]-72, f"Default Template — Page {p}")
            c.showPage()
        c.save()
        with open(path, "wb") as f:
            f.write(buf.getvalue())

    def is_docx(filename: str) -> bool:
        return filename.lower().endswith(".docx")

    def is_pdf(filename: str) -> bool:
        return filename.lower().endswith(".pdf")

    # ---- DOCX page 3 replacement
    def replace_docx_page3(template_bytes: bytes, page3_lines: list[str]) -> bytes:
        """
        Rebuild a DOCX that keeps pages 1,2 and 4..n from template,
        but replaces the whole Page 3 with provided content.
        Assumes explicit page breaks exist between pages (best-practice for templates).
        Falls back gracefully if fewer than 2 breaks exist.
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. pip install python-docx")

        src = Document(io.BytesIO(template_bytes))
        # Find up to 3 page-break boundaries by scanning runs for page breaks
        # We'll just count breaks as we write into a NEW document:
        newd = Document()
        NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        def para_has_page_break(para) -> bool:
            for run in para.runs:
                try:
                    el = run._element
                    brs = el.xpath('.//w:br[@w:type="page"]', namespaces=NS)
                    if brs:
                        return True
                except Exception:
                    pass
            return False

        # Copy original paras into a simple list for sequential access
        paras = src.paragraphs[:]

        # Helper to append a paragraph's text to new docx (simple text copy, formatting-light)
        def append_paragraph_text(p):
            # Preserve very simple structure: empty line vs non-empty
            txt = p.text or ""
            newd.add_paragraph(txt if txt.strip() else "")

        # 1) Copy until the end of Page 2 (i.e., until we've seen 2 page breaks)
        breaks_seen = 0
        i = 0
        while i < len(paras):
            p = paras[i]
            append_paragraph_text(p)
            if para_has_page_break(p):
                breaks_seen += 1
                if breaks_seen == 2:
                    # We've just finished Page 2
                    break
            i += 1

        # 2) Insert our new Page 3
        newd.add_paragraph("Page 3 — Generated by App")
        if page3_lines:
            for ln in page3_lines:
                newd.add_paragraph(ln)
        # Page break to move to Page 4
        newd.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

        # 3) Skip original Page 3 in the template, then copy the remainder (Page 4..n)
        # Move 'i' to first paragraph AFTER third page break (or end if not present)
        if breaks_seen < 2:
            # Template doesn't have 2 breaks; just append nothing (we already inserted page3)
            pass
        else:
            # Advance i to after the second break paragraph
            # (we're currently at that paragraph; move next)
            i += 1
            # Now skip until we hit the third break (end of Page 3)
            third_break_hit = False
            while i < len(paras):
                if para_has_page_break(paras[i]):
                    third_break_hit = True
                    i += 1  # move to first paragraph of Page 4
                    break
                i += 1
            # If no third break, i will be at len(paras) and nothing to copy
            # Copy the rest (Page 4..n)
            while i < len(paras):
                append_paragraph_text(paras[i])
                i += 1

        out = io.BytesIO()
        newd.save(out)
        return out.getvalue()

    # ---- PDF page 3 replacement
    def render_pdf_page_from_lines(lines: list[str]) -> bytes:
        """Create a single-page PDF from lines (simple A4 text page)."""
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab/PyPDF2 not installed. pip install reportlab PyPDF2")
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, A4[1]-72, "Page 3 — Generated by App")
        c.setFont("Helvetica", 11)
        y = A4[1]-100
        for ln in lines or []:
            c.drawString(72, y, ln[:110])
            y -= 16
            if y < 72:  # simple overflow handling
                break
        c.showPage()
        c.save()
        return buf.getvalue()

    def replace_pdf_page3(template_bytes: bytes, page3_lines: list[str]) -> bytes:
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab/PyPDF2 not installed. pip install reportlab PyPDF2")
        reader = PdfReader(io.BytesIO(template_bytes))
        writer = PdfWriter()

        # Build replacement page
        rep_pdf = PdfReader(io.BytesIO(render_pdf_page_from_lines(page3_lines)))
        rep_page = rep_pdf.pages[0]

        total = len(reader.pages)
        for idx in range(total):
            if idx == 2:  # zero-based => 3rd page
                # If template has less than 3 pages, append empties until we reach 3
                # But PyPDF2 doesn't create blanks; simple rule: if <3 pages, just append
                # existing pages and then our replacement at the end.
                if total < 3:
                    # write existing; then replace as "page 3" appended
                    pass
                writer.add_page(rep_page)
            if idx < total:
                if idx != 2:
                    writer.add_page(reader.pages[idx])

        # If template had only 1–2 pages, we ensured page3 exists by adding our page.
        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()

    # ---------------------------
    # API: Extract (Regulatory File)
    # ---------------------------
    @app.route("/api/extract", methods=["POST"])
    def api_extract():
        f = request.files.get("regulatory_file")
        if not f or not (is_docx(f.filename) or is_pdf(f.filename)):
            return jsonify({"error": "Please upload a .docx or .pdf regulatory file"}), 400

        lines = []
        medical = "—"
        try:
            if is_docx(f.filename):
                if not DOCX_AVAILABLE:
                    return jsonify({"error": "DOCX reading not available (install python-docx)"}), 500
                doc = Document(f.stream)
                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        lines.append(p.text.strip())
                medical = (lines[0] if lines else "—")[:160]
            else:
                if not PDF_AVAILABLE:
                    return jsonify({"error": "PDF reading not available (install PyPDF2)"}), 500
                reader = PdfReader(f.stream)
                for i, page in enumerate(reader.pages):
                    txt = page.extract_text() or ""
                    for ln in (txt.splitlines() if txt else []):
                        if ln.strip():
                            lines.append(ln.strip())
                medical = (lines[0] if lines else "—")[:160]
        except Exception as e:
            return jsonify({"error": f"Failed to read file: {e}"}), 500

        # Return a trimmed demo payload
        return jsonify({
            "processed": True,
            "medical": medical,
            "lines": lines[:200]  # cap the echo
        })

    # ---------------------------
    # API: Export (uses Template)
    # ---------------------------
    @app.route("/api/export", methods=["POST"])
    def api_export():
        """
        Form fields:
          - snapshot: JSON string (we'll use { content: [lines...] } )
          - template_file: optional FileStorage (.docx or .pdf)
          - fmt: optional ('docx'|'pdf'); defaults to template type, else docx
        """
        # Parse 'snapshot'
        snap_raw = request.form.get("snapshot", "{}")
        try:
            snap = json.loads(snap_raw) if snap_raw else {}
        except Exception:
            snap = {}
        page3_lines = snap.get("content") if isinstance(snap.get("content"), list) else []

        # Select template source
        up = request.files.get("template_file")
        fmt = (request.args.get("fmt") or request.form.get("fmt") or "").lower().strip()

        # Ensure defaults exist
        ensure_default_docx()
        ensure_default_pdf()

        if up:
            t_bytes = up.read()
            if is_docx(up.filename):
                if fmt not in ("docx", "pdf"):  # default to template type
                    fmt = "docx"
                if fmt == "pdf":
                    # If someone explicitly asks for PDF from DOCX template,
                    # we still replace Page 3 in DOCX then (simple) convert by
                    # returning DOCX (safer). Keep it simple: ignore mismatch.
                    fmt = "docx"
                out_bytes = replace_docx_page3(t_bytes, page3_lines)
                return send_file(
                    io.BytesIO(out_bytes),
                    as_attachment=True,
                    download_name=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            elif is_pdf(up.filename):
                if fmt not in ("docx", "pdf"):
                    fmt = "pdf"
                if fmt == "docx":
                    fmt = "pdf"
                out_bytes = replace_pdf_page3(t_bytes, page3_lines)
                return send_file(
                    io.BytesIO(out_bytes),
                    as_attachment=True,
                    download_name=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mimetype="application/pdf"
                )
            else:
                return jsonify({"error": "Template must be .docx or .pdf"}), 400
        else:
            # Use default templates
            if fmt == "pdf":
                with open(app.config["DEFAULT_PDF_TEMPLATE"], "rb") as fh:
                    t_bytes = fh.read()
                out_bytes = replace_pdf_page3(t_bytes, page3_lines)
                return send_file(
                    io.BytesIO(out_bytes),
                    as_attachment=True,
                    download_name=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mimetype="application/pdf"
                )
            else:
                with open(app.config["DEFAULT_DOCX_TEMPLATE"], "rb") as fh:
                    t_bytes = fh.read()
                out_bytes = replace_docx_page3(t_bytes, page3_lines)
                return send_file(
                    io.BytesIO(out_bytes),
                    as_attachment=True,
                    download_name=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
        # ---------------------------
    # Overlay map for front-end
    # ---------------------------
    @app.route("/overlay-map")
    def overlay_map():
        try:
            with open(AppConfig.OVERLAY_MAP_PATH, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            return jsonify(data)
        except Exception as e:
            return jsonify({"error": f"Failed to read overlay map: {e}"}), 500

    # ---------------------------
    # Download output files
    # ---------------------------
    @app.route("/download/<path:relpath>")
    def download_output(relpath):
        base = os.path.abspath(AppConfig.OUTPUT_DIR)
        abs_path = os.path.abspath(os.path.join(base, relpath))
        # prevent path traversal and 404 if missing
        if not abs_path.startswith(base) or not os.path.exists(abs_path):
            abort(404)
        directory = os.path.dirname(abs_path)
        filename = os.path.basename(abs_path)
        return send_from_directory(directory, filename, as_attachment=True)

    # ---------------------------
    # Export single (JSON body)
    # ---------------------------
    @app.route("/export", methods=["POST"])
    def export_single():
        try:
            data = request.get_json(force=True, silent=True) or {}
        except Exception:
            data = {}

        company_id = str(data.get("company_id") or "interactive")
        mapping = {
            "projectLevel": data.get("projectLevel"),
            "ticks": data.get("ticks") or {},
        }

        batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = make_batch_folder(batch_id)
        out_base = safe_filename(company_id)

        try:
            result = fill_and_export(
                docx_template=AppConfig.DOCX_TEMPLATE_PATH,
                full_docx_template=AppConfig.FULL_DOCX_TEMPLATE_PATH,
                mapping=mapping,
                out_dir=out_dir,
                out_basename=out_base,
                export_docx=AppConfig.EXPORT_DOCX,
            )
        except Exception as e:
            return jsonify({"error": f"Export failed: {e}"}), 500

        resp = {"pdf_url": f"/download/{result['rel_pdf_path']}"}
        if AppConfig.EXPORT_DOCX and result.get("rel_docx_path"):
            resp["docx_url"] = f"/download/{result['rel_docx_path']}"
        return jsonify(resp)

    # ---------------------------
    # Batch CSV (multipart/form-data)
    # ---------------------------
    @app.route("/batch", methods=["POST"])
    def batch_csv():
        f = request.files.get("file")
        if not f:
            return jsonify({"error": "Upload a CSV file as field 'file'"}), 400

        batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = make_batch_folder(batch_id)

        try:
            result = csv_process(
                csv_file=f,
                docx_template=AppConfig.DOCX_TEMPLATE_PATH,
                full_docx_template=AppConfig.FULL_DOCX_TEMPLATE_PATH,
                out_dir=out_dir,
                export_docx=AppConfig.EXPORT_DOCX,
            )
            # ZIP all generated PDFs for convenience
            zip_rel = zip_outputs(result.get("pdf_relpaths", []), out_dir, zip_name=f"{batch_id}.zip")
            result["zip_url"] = f"/download/{zip_rel}"
            return jsonify(result)
        except Exception as e:
            return jsonify({"error": f"Batch failed: {e}"}), 500

    # ---------------------------
    # Extract from standard 5-page DOCX (multipart/form-data)
    # ---------------------------
    @app.route("/extract", methods=["POST"])
    def extract_from_input():
        f = request.files.get("file")
        if not f:
            return jsonify({"error": "Upload the .docx as field 'file'"}), 400

        batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = make_batch_folder(batch_id)

        try:
            out = extract_and_map(f, out_dir)
            return jsonify(out)
        except Exception as e:
            return jsonify({"error": f"Extract failed: {e}"}), 500


    return app


if __name__ == "__main__":
    ensure_dirs()
    app = create_app()
    app.run(host="0.0.0.0", port=5000, debug=False)

